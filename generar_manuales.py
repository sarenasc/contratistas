"""
Genera los dos manuales Word del Sistema de Gestión de Contratistas:
  - Manual_Usuario.docx
  - Manual_Programador.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

HOY = datetime.date.today().strftime("%d/%m/%Y")
VERSION = "1.0"

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return p

def para(doc, text, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold   = bold
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    return p

def bullet(doc, items, style='List Bullet'):
    for item in items:
        doc.add_paragraph(item, style=style)

def tabla(doc, headers, rows, header_color='1F3864'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Cabecera
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, header_color)
    # Filas
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            t.rows[ri + 1].cells[ci].text = str(val)
    return t

def portada(doc, titulo, subtitulo):
    doc.add_paragraph()
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(titulo)
    r.font.size  = Pt(24)
    r.bold       = True
    r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(subtitulo)
    r2.font.size  = Pt(14)
    r2.italic     = True

    doc.add_paragraph()
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Sistema de Gestión de Personal Contratista\nCondor de Apalta (Ex Almahue)").font.size = Pt(12)

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run(f"Versión {VERSION}  —  {HOY}").font.size = Pt(10)
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
#  MANUAL DE USUARIO
# ══════════════════════════════════════════════════════════════════════════════

def crear_manual_usuario():
    doc = Document()

    # Márgenes
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3)
        sec.right_margin  = Cm(2.5)

    portada(doc,
            "Manual de Usuario",
            "Sistema de Gestión de Personal Contratista")

    # ── 1. INTRODUCCIÓN ───────────────────────────────────────────────────────
    heading(doc, "1. Introducción")
    para(doc,
         "Este manual describe el uso del Sistema de Gestión de Personal Contratista de Condor de Apalta "
         "(Ex Almahue). El sistema permite administrar la dotación de personal contratista, cargar asistencia "
         "diaria, gestionar tarifas y generar pre-facturas semanales a través de un flujo de aprobación "
         "multi-nivel.")
    doc.add_paragraph()

    # ── 2. ACCESO AL SISTEMA ──────────────────────────────────────────────────
    heading(doc, "2. Acceso al Sistema")
    heading(doc, "2.1  Primer Inicio (Configuración)", level=2)
    para(doc,
         "La primera vez que se accede al sistema, antes de que existan usuarios, se mostrará "
         "automáticamente la pantalla de Configuración Inicial. En ella se deben completar dos pasos:")
    bullet(doc, [
        "Paso 1 — Conexión a Base de Datos: ingresar el servidor SQL Server, usuario, contraseña "
        "y nombre de la base de datos (Fact_contratista).",
        "Paso 2 — Crear Usuario Administrador: nombre, apellido, email, usuario de login y contraseña "
        "(mínimo 6 caracteres).",
        "Presionar 'Configurar Sistema'. Si todo es correcto se redirige al login automáticamente.",
    ])

    heading(doc, "2.2  Inicio de Sesión", level=2)
    para(doc,
         "Ingresar a la URL del sistema. Completar usuario y contraseña y presionar 'Iniciar Sesión'. "
         "Si las credenciales son correctas, el sistema redirige a la pantalla de Inicio.")
    para(doc,
         "En caso de contraseña incorrecta se muestra el mensaje 'Usuario o contraseña incorrectos'. "
         "El sistema no bloquea la cuenta, pero la contraseña siempre se valida de forma segura.")

    heading(doc, "2.3  Pantalla de Inicio", level=2)
    para(doc,
         "Al ingresar se muestra la pantalla de bienvenida con el nombre del usuario. "
         "Si el usuario tiene perfil de Aprobador o Administrador, aparecerán alertas automáticas en caso de:")
    bullet(doc, [
        "Lotes de asistencia pendientes de aprobación (fondo amarillo).",
        "Lotes rechazados que requieren atención (fondo rojo).",
    ])
    para(doc,
         "Cada alerta incluye un botón directo a la bandeja correspondiente.")

    heading(doc, "2.4  Cerrar Sesión", level=2)
    para(doc, "Hacer clic en 'Cerrar sesión' en la parte superior derecha del menú de navegación.")
    doc.add_paragraph()

    # ── 3. NAVEGACIÓN Y PERFILES ──────────────────────────────────────────────
    heading(doc, "3. Navegación y Perfiles de Usuario")
    para(doc,
         "El menú superior muestra sólo los módulos a los que el usuario tiene acceso según su perfil. "
         "Existen cuatro perfiles:")
    tabla(doc,
          ["Perfil", "Acceso"],
          [
              ["Administrador",     "Acceso total a todos los módulos y configuraciones."],
              ["Edición",           "Puede crear y editar registros. No puede aprobar asistencia."],
              ["Aprobador-Edición", "Puede editar y además aprobar o rechazar lotes de asistencia."],
              ["Visualización",     "Solo puede ver información. No puede modificar nada."],
          ])
    doc.add_paragraph()
    para(doc,
         "Adicionalmente, cada usuario puede tener acceso sólo a ciertos módulos: "
         "Configuraciones, Tarifas, Procesos, Contratistas, Aprobación.")
    doc.add_paragraph()

    # ── 4. MÓDULO CONFIGURACIONES ─────────────────────────────────────────────
    heading(doc, "4. Módulo Configuraciones")
    para(doc,
         "Disponible para usuarios con acceso al módulo 'configuraciones'. "
         "Contiene los catálogos base del sistema.")

    heading(doc, "4.1  Tipo de Mano de Obra", level=2)
    para(doc, "Permite crear y editar los tipos de mano de obra (ej: Directa, Indirecta) con su abreviatura.")

    heading(doc, "4.2  Labores (Cargos)", level=2)
    para(doc,
         "Catálogo de labores o cargos que realizan los trabajadores. Permite agregar uno a uno "
         "o mediante carga masiva desde un archivo Excel (hoja 'TIPO CARGO').")
    bullet(doc, [
        "Columnas requeridas en Excel: CARGO, TIPO MO (número: 1 o 2), COD FACT.",
        "Se pueden repetir cargos si corresponde a distintos contratos.",
    ])

    heading(doc, "4.3  Áreas", level=2)
    para(doc, "Gestión de las áreas de trabajo (ej: Packing, Bodega, Clasificado). Cada área tiene un código de facturación.")

    heading(doc, "4.4  Turnos", level=2)
    para(doc, "Define los turnos de trabajo (ej: Turno 1, Turno 2, Turno Noche).")

    heading(doc, "4.5  Jefes de Área", level=2)
    para(doc,
         "Registro de los jefes responsables por área y turno. "
         "Un área puede tener más de un jefe (uno por turno). "
         "Se puede vincular cada jefe a su usuario de sistema en el campo 'Usuario sistema'.")
    bullet(doc, [
        "Si un jefe cubre ambos turnos, dejar el campo Turno vacío.",
        "El campo 'Nivel de Aprobación' distingue entre Jefe de Área (1) y Jefe de Operaciones (2).",
    ])

    heading(doc, "4.6  Usuarios", level=2)
    para(doc,
         "Exclusivo para Administradores. Permite crear, editar y activar/desactivar usuarios del sistema.")
    bullet(doc, [
        "Datos básicos: nombre, apellido, email, usuario de login, contraseña, área propia.",
        "Perfil: define el nivel de acceso general.",
        "Módulos: checkboxes para cada módulo al que tendrá acceso.",
        "Áreas de aprobación: multiselect con las áreas que el usuario puede aprobar.",
        "Cargos específicos: para casos donde el usuario aprueba cargos de otras áreas "
        "(ej: un jefe de Packing que también aprueba 'Armados de Materiales' de Bodega).",
        "Para cambiar la contraseña, ingresar la nueva en el campo correspondiente al editar. "
        "Si se deja vacío, la contraseña no cambia.",
    ])
    doc.add_paragraph()

    # ── 5. MÓDULO TARIFAS ─────────────────────────────────────────────────────
    heading(doc, "5. Módulo Tarifas")

    heading(doc, "5.1  Tipos de Tarifa", level=2)
    para(doc,
         "Define las tarifas regulares que se aplican a cada labor. Cada tarifa tiene:")
    bullet(doc, [
        "Valor Contratista: monto diario base.",
        "Horas Extras: valor por hora extra.",
        "% Contratista (Jornada) y % HH.EE: porcentaje de recargo que paga el empleador.",
        "Bono: monto fijo adicional por jornada.",
        "Kilo / Caja: indica si la tarifa se calcula por unidad física.",
        "Fecha desde / hasta y estado Activo.",
    ])
    para(doc, "Solo puede haber una tarifa activa por período. El sistema valida que no se superpongan fechas.")

    heading(doc, "5.2  Tarifas Especiales", level=2)
    para(doc,
         "Tarifas para días especiales (feriados, eventos). Incluyen valor base y porcentajes de recargo.")

    heading(doc, "5.3  Cargos y Tarifas", level=2)
    para(doc,
         "Vincula cada labor (cargo) con su tipo de tarifa y especie. "
         "También permite carga masiva desde Excel (hoja 'Tarifas') con columnas: Labor | Tipo Tarifa | Especie.")
    doc.add_paragraph()

    # ── 6. MÓDULO CONTRATISTAS ────────────────────────────────────────────────
    heading(doc, "6. Módulo Contratistas")

    heading(doc, "6.1  Contratistas", level=2)
    para(doc, "Gestión del catálogo de empresas contratistas con RUT, razón social y código de facturación.")

    heading(doc, "6.2  Solicitud de Contratistas", level=2)
    para(doc,
         "Registro de solicitudes de dotación por parte del área hacia el contratista. "
         "Cada solicitud indica: contratista, labor, área, turno, cantidad solicitada y fecha.")
    doc.add_paragraph()

    # ── 7. MÓDULO PROCESOS ────────────────────────────────────────────────────
    heading(doc, "7. Módulo Procesos")

    heading(doc, "7.1  Carga de Asistencia", level=2)
    para(doc,
         "Proceso principal del sistema. Se ejecuta en dos pasos:")
    heading(doc, "Paso 1 — Subir archivo Excel", level=3)
    bullet(doc, [
        "Seleccionar el archivo Excel de asistencia (formato del reloj biométrico).",
        "Opcionalmente, aplicar un filtro por semana, día específico o cargar todo.",
        "El sistema lee el archivo y detecta los valores únicos de Área, Empleador, Cargo y Turno.",
    ])
    heading(doc, "Paso 2 — Mapear y cargar", level=3)
    bullet(doc, [
        "Para cada valor detectado en el Excel, seleccionar el registro correspondiente en el sistema.",
        "El sistema recuerda los mapeos previos para facilitar futuras cargas.",
        "Presionar 'Iniciar Carga'. Se muestra una barra de progreso mientras se insertan los registros.",
        "Al finalizar, el lote queda en estado Pendiente y pasa al flujo de aprobación.",
    ])

    heading(doc, "7.2  Editar Asistencia", level=2)
    para(doc,
         "Permite al personal de RRHH corregir registros de asistencia ya cargados sin necesidad "
         "de volver a subir el Excel completo.")
    bullet(doc, [
        "Seleccionar el lote a editar desde el selector desplegable.",
        "Editar directamente en la tabla: Jornada, HH.EE, Labor u Observación por fila.",
        "Al guardar, si el lote estaba en estado Rechazado, vuelve automáticamente a Pendiente "
        "para que el jefe correspondiente vuelva a revisarlo.",
    ])

    heading(doc, "7.3  Descuento", level=2)
    para(doc, "Permite registrar descuentos por contratista que se aplican en la Pre-Factura.")

    heading(doc, "7.4  Pre-Factura", level=2)
    para(doc,
         "Genera el resumen de facturación por semana y contratista.")
    bullet(doc, [
        "IMPORTANTE: Solo se puede calcular la Pre-Factura cuando TODOS los lotes de "
        "la semana/año seleccionada tienen estado 'listo_factura'.",
        "Si hay lotes pendientes de aprobación, el sistema muestra cuáles son y bloquea el cálculo.",
        "Muestra el desglose por labor: base jornada, base HH.EE, porcentaje contratista y total.",
    ])

    heading(doc, "7.5  Proformas", level=2)
    para(doc,
         "Gestión de las proformas generadas. Permite revisar, cerrar o exportar el detalle.")
    doc.add_paragraph()

    # ── 8. MÓDULO APROBACIÓN ──────────────────────────────────────────────────
    heading(doc, "8. Módulo Aprobación")
    para(doc,
         "Flujo de aprobación de la asistencia cargada. Consta de dos etapas secuenciales "
         "antes de que la asistencia quede disponible para Pre-Factura.")

    heading(doc, "8.1  Flujo completo de aprobación", level=2)
    para(doc, "El flujo tiene los siguientes estados:")
    tabla(doc,
          ["Estado", "Descripción", "Quién actúa"],
          [
              ["pendiente",      "Lote recién cargado por RRHH.",                                           "Jefe de Área"],
              ["aprobado_area",  "Todos los jefes de área aprobaron su parte.",                             "Jefe de Operaciones"],
              ["rechazado_area", "Al menos un jefe de área rechazó. RRHH debe editar y reenviar.",          "RRHH"],
              ["rechazado_ops",  "Jefe de Operaciones rechazó. Vuelve a jefes de área y luego a RRHH.",     "Jefe de Área → RRHH"],
              ["listo_factura",  "Aprobación completa. Disponible para Pre-Factura.",                       "—"],
          ])
    doc.add_paragraph()

    heading(doc, "8.2  Bandeja Jefe de Área", level=2)
    para(doc,
         "Vista disponible para usuarios con perfil Aprobador-Edición. "
         "Muestra los lotes en estado 'pendiente' o 'rechazado_ops' que contienen "
         "registros de las áreas asignadas al jefe.")
    bullet(doc, [
        "Ver: abre el detalle completo del lote.",
        "Aprobar: registra la aprobación. Si todos los jefes requeridos aprobaron, el lote pasa a 'aprobado_area'.",
        "Rechazar: debe ingresar una observación obligatoria explicando el problema. "
        "El lote pasa a 'rechazado_area' y RRHH puede editarlo.",
    ])

    heading(doc, "8.3  Bandeja Jefe de Operaciones", level=2)
    para(doc,
         "Vista para el Jefe de Operaciones (nivel de aprobación 2). "
         "Muestra los lotes en estado 'aprobado_area' listos para revisión final.")
    bullet(doc, [
        "Muestra el historial de aprobaciones de los jefes de área para cada lote.",
        "Aprobar: el lote pasa a 'listo_factura' y queda disponible para Pre-Factura.",
        "Rechazar: debe indicar el área con el problema y una observación. "
        "El lote vuelve a 'rechazado_ops' y los jefes de área deben revisarlo nuevamente.",
    ])

    heading(doc, "8.4  Detalle de Asistencia", level=2)
    para(doc,
         "Vista de solo lectura con todos los registros del lote seleccionado. "
         "Incluye totales de jornada y HH.EE, y el historial completo de aprobaciones con fechas y observaciones.")
    doc.add_paragraph()

    # ── 9. FLUJO GENERAL DEL SISTEMA ─────────────────────────────────────────
    heading(doc, "9. Flujo General del Sistema")
    para(doc, "El proceso completo semanal sigue este orden:")
    bullet(doc, [
        "1. RRHH sube el archivo Excel de asistencia (Procesos → Carga Asistencia).",
        "2. El lote queda en estado 'pendiente'.",
        "3. Los Jefes de Área ingresan a su bandeja y aprueban o rechazan.",
        "   → Si aprueban todos: el lote pasa a 'aprobado_area'.",
        "   → Si alguno rechaza: RRHH recibe la observación, edita y el lote vuelve a 'pendiente'.",
        "4. El Jefe de Operaciones revisa el lote 'aprobado_area'.",
        "   → Si aprueba: lote queda 'listo_factura'.",
        "   → Si rechaza: el lote vuelve a 'rechazado_ops' → Jefes de Área → RRHH → ciclo.",
        "5. Con todos los lotes de la semana en 'listo_factura', RRHH genera la Pre-Factura.",
        "6. Se revisa y cierra la proforma correspondiente.",
    ])
    doc.add_paragraph()

    # ── 10. PREGUNTAS FRECUENTES ──────────────────────────────────────────────
    heading(doc, "10. Preguntas Frecuentes")
    preguntas = [
        ("¿Qué pasa si me equivoco al cargar la asistencia?",
         "Usar 'Editar Asistencia' para corregir los datos. Si el lote ya fue aprobado por el jefe, "
         "al editar vuelve a estado 'pendiente' automáticamente y el jefe debe volver a aprobar."),
        ("¿Puedo cargar asistencia de varios días en un solo archivo?",
         "Sí. En el Paso 1 puede seleccionar el filtro 'Todo' para cargar todos los días del archivo, "
         "o filtrar por semana o día específico."),
        ("¿Por qué no puedo generar la Pre-Factura?",
         "El sistema bloquea la Pre-Factura si existen lotes de esa semana sin aprobación final. "
         "Revisar la alerta que indica qué lotes faltan y en qué estado están."),
        ("¿Qué ocurre si un jefe de área no está disponible para aprobar?",
         "Si el área tiene más de un jefe (un jefe por turno), el otro puede cubrir ambos. "
         "El administrador puede asignar ambos turnos al jefe disponible desde Configuraciones → Jefes de Área."),
        ("¿Cómo cambio mi contraseña?",
         "Contactar al Administrador del sistema para que la actualice desde Configuraciones → Usuarios."),
    ]
    for preg, resp in preguntas:
        p = doc.add_paragraph()
        r = p.add_run(f"• {preg}")
        r.bold = True
        doc.add_paragraph(f"  {resp}")
        doc.add_paragraph()

    ruta = r"C:\xampp\htdocs\contratista\Manual_Usuario.docx"
    doc.save(ruta)
    print(f"OK Guardado: {ruta}")


# ══════════════════════════════════════════════════════════════════════════════
#  MANUAL DEL PROGRAMADOR
# ══════════════════════════════════════════════════════════════════════════════

def crear_manual_programador():
    doc = Document()

    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3)
        sec.right_margin  = Cm(2.5)

    portada(doc,
            "Manual del Programador",
            "Sistema de Gestión de Personal Contratista")

    # ── 1. VISIÓN GENERAL ─────────────────────────────────────────────────────
    heading(doc, "1. Visión General del Sistema")
    para(doc,
         "El Sistema de Gestión de Personal Contratista es una aplicación web desarrollada en PHP 8 "
         "con SQL Server como motor de base de datos. Corre sobre XAMPP en Windows Server. "
         "No utiliza frameworks MVC; sigue una arquitectura de scripts PHP por página con helpers compartidos.")

    tabla(doc,
          ["Tecnología", "Versión / Detalle"],
          [
              ["PHP",           "8.x (con extensión sqlsrv)"],
              ["SQL Server",    "2019+ (driver sqlsrv para PHP)"],
              ["Servidor Web",  "Apache (XAMPP)"],
              ["Frontend",      "Bootstrap 5.3, jQuery 3.7, Select2 4.1"],
              ["Excel",         "PhpSpreadsheet (via Composer)"],
              ["Sistema Oper.", "Windows Server / Windows 11"],
          ])
    doc.add_paragraph()

    # ── 2. ESTRUCTURA DE DIRECTORIOS ──────────────────────────────────────────
    heading(doc, "2. Estructura de Directorios")
    estructura = """\
C:\\xampp\\htdocs\\contratista\\
│
├── app\\
│   ├── lib\\
│   │   ├── db.php           ← función db_query() con manejo de errores
│   │   └── permisos.php     ← helpers de control de acceso
│   └── Middlewares\\
│       └── auth_guard.php   ← redirige al login si no hay sesión
│
├── config\\
│   ├── app.php              ← define BASE_URL
│   ├── conexion.php         ← abre $conn (y $conn2 opcional)
│   ├── .env                 ← credenciales BD (no versionar)
│   └── setup.lock           ← existe tras el primer setup
│
├── public\\
│   ├── _bootstrap.php       ← punto de entrada de todas las páginas
│   ├── index.php            ← pantalla de login
│   ├── login.php            ← procesa el POST del login
│   ├── logout.php           ← destruye la sesión
│   ├── setup.php            ← configuración inicial (solo si no hay lock)
│   ├── Inicio.php           ← pantalla de bienvenida con notificaciones
│   │
│   ├── aprobacion\\
│   │   ├── bandeja_jefe.php
│   │   ├── bandeja_operaciones.php
│   │   └── detalle_asistencia.php
│   │
│   ├── configuraciones\\
│   │   ├── Area.php
│   │   ├── Cargos.php
│   │   ├── JefeArea.php
│   │   ├── Reg_Usuario.php
│   │   ├── Turnos.php
│   │   └── tipo_mo.php
│   │
│   ├── procesos\\
│   │   ├── carga_asistencia.php          ← UI pasos 1 y 2
│   │   ├── carga_asistencia_ajax_start.php
│   │   ├── carga_asistencia_ajax_chunk.php
│   │   ├── carga_asistencia_paso2_start.php
│   │   ├── carga_asistencia_paso2_chunk.php
│   │   ├── editar_asistencia.php
│   │   ├── Pre-Factura.php
│   │   ├── proformas.php
│   │   ├── descuento.php
│   │   └── guardar_factura.php
│   │
│   ├── tarifas\\
│   │   ├── tipo_tarifa.php
│   │   ├── tarifasEspecial.php
│   │   └── Tarifas_Cargo.php
│   │
│   ├── contratista\\
│   │   ├── ingreso_contratista.php
│   │   └── Solicitud_Contra.php
│   │
│   ├── partials\\
│   │   ├── header.php
│   │   ├── footer.php
│   │   ├── navbar.php       ← ítems condicionales por perfil
│   │   ├── navbar_wrapper.php
│   │   └── base_url.php
│   │
│   └── assets\\
│       └── css\\app.css
│
├── storage\\asistencia\\     ← archivos Excel subidos temporalmente
├── vendor\\                  ← Composer (PhpSpreadsheet)
├── setup.sql                 ← script completo de BD
└── Plan_Perfiles_Usuarios.md ← documentación de diseño
"""
    doc.add_paragraph(estructura).style = 'No Spacing'
    doc.add_paragraph()

    # ── 3. BASE DE DATOS ──────────────────────────────────────────────────────
    heading(doc, "3. Base de Datos")
    heading(doc, "3.1  Nombre y Motor", level=2)
    para(doc,
         "Base de datos: Fact_contratista — SQL Server. "
         "La base anterior (proforma_contratista) se mantiene como respaldo en contratista2/ y NO se modifica.")

    heading(doc, "3.2  Tablas Principales", level=2)
    tabla(doc,
          ["Tabla", "Descripción"],
          [
              ["dbo.Area",                      "Áreas de trabajo con código de facturación."],
              ["dbo.dota_perfiles",             "4 perfiles de usuario: Admin, Edición, Aprobador-Edición, Visualización."],
              ["dbo.dota_usuarios",             "Usuarios del sistema (reemplaza TRA_usuario). Incluye perfil, área propia, email y activo."],
              ["dbo.dota_usuario_modulos",      "Módulos accesibles por usuario (configuraciones, tarifas, procesos, contratista, aprobacion)."],
              ["dbo.dota_usuario_areas",        "Áreas que un usuario aprobador puede aprobar."],
              ["dbo.dota_usuario_cargos",       "Cargos específicos fuera del área propia que puede aprobar un usuario."],
              ["dbo.dota_tipo_mo",              "Tipos de mano de obra (Directa, Indirecta)."],
              ["dbo.Dota_Cargo",                "Labores/cargos de los trabajadores."],
              ["dbo.especie",                   "Catálogo de especies (fruta)."],
              ["dbo.dota_contratista",          "Empresas contratistas."],
              ["dbo.Dota_tipo_tarifa",          "Tarifas regulares por período con porcentajes de recargo."],
              ["dbo.Dota_Valor_Dotacion",       "Vincula cargo + tarifa + especie (+ contratista opcional)."],
              ["dbo.Dota_Tarifa_Especiales",    "Tarifas para días especiales (feriados)."],
              ["dbo.Dota_ValorEspecial_Dotacion","Valores especiales por cargo + fecha."],
              ["dbo.dota_turno",                "Turnos de trabajo."],
              ["dbo.dota_jefe_area",            "Jefes de área con su área, turno, nivel de aprobación y vínculo a id_usuario."],
              ["dbo.dota_asistencia_lote",      "Un registro por archivo Excel cargado. Estado: pendiente → listo_factura."],
              ["dbo.dota_asistencia_carga",     "Filas individuales de asistencia (un registro por trabajador × día)."],
              ["dbo.dota_asistencia_aprobacion","Log de aprobaciones, rechazos y ediciones por usuario y lote."],
              ["dbo.dota_asistencia_mapa",      "Memoriza mapeos Excel → sistema (área, empleador, cargo, turno)."],
              ["dbo.dota_solicitud_contratista","Solicitudes de dotación: contratista, cargo, área, turno, cantidad."],
              ["dbo.dota_descuento",            "Descuentos por contratista aplicables en la pre-factura."],
              ["dbo.dota_factura",              "Cabecera de proformas (semana, año, versión, estado)."],
              ["dbo.dota_factura_detalle",      "Detalle por contratista y labor de cada proforma."],
              ["dbo.dota_factura_descuento",    "Descuentos específicos dentro de una proforma."],
          ])
    doc.add_paragraph()

    heading(doc, "3.3  Vista Principal", level=2)
    para(doc,
         "dbo.vw_proformas_detalle: vista que une asistencia + tarifas (regulares y especiales) + dimensiones. "
         "Aplica prioridad de tarifa especial sobre regular. Solo incluye registros de lotes con estado = 'listo_factura'. "
         "Calcula base_jorn, base_hhee, pct_jorn, pct_hhee y total_factura.")

    heading(doc, "3.4  Estados del Lote (Máquina de Estados)", level=2)
    tabla(doc,
          ["Estado", "Siguiente(s) Estado(s)", "Trigger"],
          [
              ["pendiente",      "aprobado_area / rechazado_area",   "Jefe área aprueba o rechaza"],
              ["aprobado_area",  "listo_factura / rechazado_ops",     "Jefe operaciones aprueba o rechaza"],
              ["rechazado_area", "pendiente",                         "RRHH edita el lote"],
              ["rechazado_ops",  "pendiente (via jefe área)",         "Jefe área revisa → RRHH edita"],
              ["listo_factura",  "—",                                 "Estado final: disponible para Pre-Factura"],
          ])
    doc.add_paragraph()

    # ── 4. AUTENTICACIÓN Y PERMISOS ───────────────────────────────────────────
    heading(doc, "4. Autenticación y Permisos")

    heading(doc, "4.1  Flujo de Login", level=2)
    bullet(doc, [
        "index.php: renderiza el formulario.",
        "login.php: recibe POST, valida usuario/contraseña contra dota_usuarios con password_verify().",
        "Si es válido, carga en sesión: id_usuario, usuario, nombre, apellido, id_perfil, "
        "modulos[], areas_aprobar[], cargos_aprobar[], nivel_aprobacion.",
        "Redirige a Inicio.php.",
        "auth_guard.php: incluido desde _bootstrap.php; redirige al login si no hay id_usuario en sesión.",
    ])

    heading(doc, "4.2  Variables de Sesión", level=2)
    tabla(doc,
          ["Variable", "Tipo", "Descripción"],
          [
              ["$_SESSION['id_usuario']",       "int",    "ID del usuario logueado."],
              ["$_SESSION['usuario']",          "string", "Login del usuario."],
              ["$_SESSION['nombre']",           "string", "Nombre de pila."],
              ["$_SESSION['apellido']",         "string", "Apellido."],
              ["$_SESSION['id_perfil']",        "int",    "1=Admin, 2=Edición, 3=Aprobador-Ed., 4=Visualiz."],
              ["$_SESSION['modulos']",          "array",  "Lista de módulos accesibles."],
              ["$_SESSION['areas_aprobar']",    "array",  "IDs de áreas que puede aprobar."],
              ["$_SESSION['cargos_aprobar']",   "array",  "IDs de cargos específicos extra."],
              ["$_SESSION['nivel_aprobacion']", "int",    "0=no jefe, 1=jefe área, 2=jefe operaciones."],
          ])

    heading(doc, "4.3  Funciones Helper (app/lib/permisos.php)", level=2)
    tabla(doc,
          ["Función", "Retorna", "Descripción"],
          [
              ["es_admin()",              "bool", "true si id_perfil == 1."],
              ["puede_aprobar()",         "bool", "true si Admin o Aprobador-Edición (perfil 3)."],
              ["puede_modulo($modulo)",   "bool", "true si Admin o el módulo está en su lista."],
              ["puede_aprobar_area($id)", "bool", "true si Admin o el id_area está en areas_aprobar."],
              ["puede_aprobar_cargo($id)","bool", "true si Admin o el id_cargo está en cargos_aprobar."],
              ["es_jefe_operaciones()",   "bool", "true si nivel_aprobacion >= 2."],
              ["nombre_usuario()",        "string","Retorna 'Nombre Apellido' o el login si no hay nombre."],
          ])
    doc.add_paragraph()
    para(doc,
         "Uso en páginas: al inicio de cada página protegida, agregar la verificación correspondiente:",
         italic=True)
    doc.add_paragraph("if (!puede_modulo('procesos')) { header('Location: '.BASE_URL.'/Inicio.php'); exit; }").style = 'No Spacing'
    doc.add_paragraph()

    # ── 5. CARGA DE ASISTENCIA ────────────────────────────────────────────────
    heading(doc, "5. Carga de Asistencia — Detalle Técnico")

    heading(doc, "5.1  Paso 1 (AJAX)", level=2)
    bullet(doc, [
        "carga_asistencia.php: formulario de upload. Al subir el archivo, llama a carga_asistencia_ajax_start.php.",
        "ajax_start.php: lee el Excel con PhpSpreadsheet, filtra filas según criterio, "
        "guarda cada fila válida como JSON en un archivo temporal (storage/asistencia/), "
        "guarda los valores únicos de Área/Empleador/Cargo/Turno en $_SESSION['asistencia_upload'].",
        "ajax_chunk.php: lee chunks del archivo temporal y los devuelve para mostrar en la UI.",
        "El archivo temporal persiste entre requests (es la fuente para el Paso 2).",
    ])

    heading(doc, "5.2  Paso 2 (AJAX chunked)", level=2)
    bullet(doc, [
        "El formulario de mapeo hace POST a carga_asistencia_paso2_start.php.",
        "paso2_start: valida mapeos, actualiza turnos de jefes que rotaron, "
        "construye jefeByAreaTurno (mapa area_id+turno_id → jefe_id), "
        "lee primer registro para capturar semana y año del lote, "
        "guarda todo en $_SESSION['asistencia_paso2_state'] y responde {ok, totalRows}.",
        "La UI hace polling a paso2_chunk.php enviando el estado en sesión.",
        "paso2_chunk: lee chunkSize=800 líneas del archivo temporal, inserta en dota_asistencia_carga, "
        "devuelve {ok, done, pct, inserted, total}.",
        "Cuando done=true: inserta el lote en dota_asistencia_lote con estado='pendiente', "
        "limpia la sesión.",
        "En caso de error: hace DELETE WHERE registro = $archivo (limpieza sin transacción entre requests).",
    ])

    heading(doc, "5.3  Mapeos persistentes", level=2)
    para(doc,
         "Los mapeos Excel→sistema se guardan en dota_asistencia_mapa usando MERGE. "
         "En futuras cargas, paso2_start preselecciona automáticamente los valores conocidos.")
    doc.add_paragraph()

    # ── 6. MÓDULO DE APROBACIÓN ───────────────────────────────────────────────
    heading(doc, "6. Módulo de Aprobación — Detalle Técnico")

    heading(doc, "6.1  Lógica 'todos aprobaron'", level=2)
    para(doc, "Al aprobar en bandeja_jefe.php, se verifica con la siguiente consulta:")
    doc.add_paragraph(
        "SELECT COUNT(*) FROM (\n"
        "    SELECT DISTINCT j.id_usuario\n"
        "    FROM dbo.dota_jefe_area j\n"
        "    WHERE j.activo=1 AND j.id_usuario IS NOT NULL AND j.nivel_aprobacion=1\n"
        "      AND j.id_area IN (SELECT DISTINCT area FROM dota_asistencia_carga WHERE registro=?)\n"
        ") req\n"
        "WHERE NOT EXISTS (\n"
        "    SELECT 1 FROM dota_asistencia_aprobacion ap\n"
        "    WHERE ap.registro=? AND ap.id_usuario=req.id_usuario AND ap.accion='aprobado'\n"
        ")"
    ).style = 'No Spacing'
    para(doc,
         "Si el resultado es 0 (ningún aprobador requerido pendiente), el lote pasa a 'aprobado_area'.",
         italic=True)
    para(doc,
         "Nota: si ningún jefe tiene id_usuario asignado para las áreas del lote, "
         "el lote queda en pendiente indefinidamente. Asegurarse de vincular id_usuario en JefeArea.")
    doc.add_paragraph()

    heading(doc, "6.2  Permisos en bandejas", level=2)
    bullet(doc, [
        "bandeja_jefe.php: accesible si puede_aprobar() o es_admin(). "
        "Filtra lotes por $_SESSION['areas_aprobar'] del usuario (Admin ve todos).",
        "bandeja_operaciones.php: accesible si es_jefe_operaciones() o es_admin().",
        "detalle_asistencia.php: accesible si puede_aprobar() o puede_modulo('procesos') o es_admin().",
    ])
    doc.add_paragraph()

    # ── 7. CÁLCULO DE PRE-FACTURA ─────────────────────────────────────────────
    heading(doc, "7. Cálculo de Pre-Factura")
    para(doc, "La lógica de cálculo está en Pre-Factura.php y se basa en la vista vw_proformas_detalle.")
    tabla(doc,
          ["Concepto", "Fórmula"],
          [
              ["Base Jornada",      "jornada × ValorContratista"],
              ["Base HH.EE",        "hhee × horasExtras"],
              ["Total Empleador J", "Base Jornada × (1 + PorcContratista)"],
              ["Total Empleador H", "Base HH.EE × (1 + porc_hhee)"],
              ["Porc. Contratista J","Total Empleador J − Base Jornada"],
              ["Porc. Contratista H","Total Empleador H − Base HH.EE"],
              ["Total Factura",     "Total Empleador J + Total Empleador H + Bono"],
          ])
    para(doc,
         "La tarifa especial tiene prioridad sobre la regular cuando existe un registro "
         "en Dota_ValorEspecial_Dotacion o Dota_Tarifa_Especiales para la fecha del registro.")
    doc.add_paragraph()

    # ── 8. GESTIÓN DE RAMAS GIT ───────────────────────────────────────────────
    heading(doc, "8. Control de Versiones — Git")
    tabla(doc,
          ["Rama", "Propósito"],
          [
              ["main",       "Rama original, queda como respaldo histórico."],
              ["master",     "Código estable y probado. Solo se actualiza cuando hay cambios significativos testeados."],
              ["desarrollo", "Rama activa de desarrollo. Todos los cambios nuevos van aquí."],
          ])
    bullet(doc, [
        "Repositorio: https://github.com/sarenasc/contratistas",
        "Flujo: desarrollar en 'desarrollo' → probar → merge a 'master' cuando esté estable.",
        "No subir config/.env (contiene contraseñas). Está en .gitignore implícitamente.",
    ])
    doc.add_paragraph()

    # ── 9. AGREGAR NUEVAS PÁGINAS ─────────────────────────────────────────────
    heading(doc, "9. Guía para Agregar Nuevas Páginas")
    para(doc, "Toda página nueva debe seguir este patrón:")
    doc.add_paragraph(
        "<?php\n"
        "require_once __DIR__ . '/../_bootstrap.php';  // sesión, permisos, $conn\n\n"
        "// Verificar acceso\n"
        "if (!puede_modulo('mi_modulo')) {\n"
        "    header('Location: ' . BASE_URL . '/Inicio.php'); exit;\n"
        "}\n\n"
        "$title = 'Mi Página';\n"
        "include __DIR__ . '/../partials/header.php';\n"
        "include __DIR__ . '/../partials/navbar_wrapper.php';\n"
        "?>\n"
        "<main class=\"container py-4\">\n"
        "    <!-- contenido -->\n"
        "</main>\n"
        "<?php include __DIR__ . '/../partials/footer.php'; ?>"
    ).style = 'No Spacing'
    doc.add_paragraph()
    bullet(doc, [
        "Usar db_query($conn, $sql, $params) para consultas — lanza RuntimeException en error.",
        "Nunca construir SQL concatenando variables del usuario. Siempre usar parámetros (?).",
        "Usar htmlspecialchars() al imprimir cualquier dato de BD en HTML.",
        "Agregar el ítem de menú en public/partials/navbar.php con el condicional puede_modulo() correspondiente.",
    ])
    doc.add_paragraph()

    # ── 10. MANTENIMIENTO ─────────────────────────────────────────────────────
    heading(doc, "10. Tareas de Mantenimiento")

    heading(doc, "10.1  Reconstruir la Base de Datos", level=2)
    bullet(doc, [
        "Ejecutar setup.sql conectado a master en SQL Server Management Studio.",
        "El script crea la BD Fact_contratista si no existe y todas las tablas con IF OBJECT_ID.",
        "Es seguro ejecutarlo múltiples veces (idempotente).",
    ])

    heading(doc, "10.2  Resetear el Setup del Sistema", level=2)
    bullet(doc, [
        "Eliminar el archivo config/setup.lock.",
        "Al acceder al sistema redirigirá automáticamente a setup.php.",
        "Útil cuando se cambia de servidor de BD o se necesita crear un nuevo admin.",
    ])

    heading(doc, "10.3  Archivos temporales de asistencia", level=2)
    bullet(doc, [
        "Los archivos Excel subidos se guardan en storage/asistencia/ con nombre único (timestamp + uniqid).",
        "No se eliminan automáticamente tras la carga. Revisar periódicamente y eliminar los antiguos.",
    ])

    heading(doc, "10.4  Notificaciones por email", level=2)
    para(doc,
         "Pendiente de implementar. El punto de integración está en carga_asistencia_paso2_chunk.php "
         "al detectar done=true, después de insertar el lote. Agregar llamada a una función de envío "
         "de email con los datos del lote y los jefes de área que deben aprobar.")
    doc.add_paragraph()

    ruta = r"C:\xampp\htdocs\contratista\Manual_Programador.docx"
    doc.save(ruta)
    print(f"OK Guardado: {ruta}")


# ── Ejecutar ──────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    crear_manual_usuario()
    crear_manual_programador()
    print("\nAmbos manuales generados correctamente.")
